"""
P.E.T.E.R. — Received-file storage helper.

Stores files the user drops / selects in the desktop app under ./received_files
and tracks which one Peter should treat as the "active file" (the one the user
is talking about — e.g. their resume). Mirrors the Spotify default-playlist
pattern (a small JSON persists the active selection).

Why a folder instead of a JSON blob: files like PDFs/DOCX are binary; we keep
the original bytes so Peter can re-read them, and the UI can list real
filenames. The active-file pointer just names the one to focus on.
"""

from __future__ import annotations

import base64
import json
import logging
import os
from pathlib import Path

_HERE = Path(os.path.dirname(os.path.abspath(__file__)))
_FILES_DIR = _HERE / "received_files"
_ACTIVE_FILE = _HERE / "active_file.json"

# Keep file/folder names safe to use on disk.
_SAFE_CHARS = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789._- ")


def _safe_filename(name: str) -> str:
    """Sanitize a filename to something safe for the filesystem."""
    name = (name or "").strip().replace("\\", "_").replace("/", "_")
    cleaned = "".join(c if c in _SAFE_CHARS else "_" for c in name)
    cleaned = cleaned.strip(" .")
    return cleaned[:120] or "unnamed.txt"


def _ensure_dir() -> None:
    """Make sure the received-files folder exists."""
    try:
        _FILES_DIR.mkdir(parents=True, exist_ok=True)
    except OSError as e:  # noqa: BLE001
        logging.warning(f"Couldn't create received-files dir: {e}")


# ──────────────────────────────────────────────────────────────
# Save / list
# ──────────────────────────────────────────────────────────────
def save_file(filename: str, data: bytes) -> dict:
    """Save a file's bytes into received_files. Returns {ok, path, name}.

    `data` may be raw bytes or a base64 string (the UI sends base64). If a
    file with the same name exists, we keep both by appending a number.
    """
    _ensure_dir()
    name = _safe_filename(filename)
    if not name:
        return {"ok": False, "error": "invalid filename"}

    # Decode base64 if that's what we got.
    raw = data
    if isinstance(data, str):
        try:
            raw = base64.b64decode(data)
        except Exception:  # noqa: BLE001
            raw = data.encode("utf-8", errors="replace")

    # Avoid overwriting an existing file — add a numeric suffix instead.
    target = _FILES_DIR / name
    base, ext = os.path.splitext(name)
    n = 1
    while target.exists():
        target = _FILES_DIR / f"{base}_{n}{ext}"
        n += 1

    try:
        target.write_bytes(raw)
        # Grab a short text preview so the UI can show confirmation.
        preview = _read_text_preview(raw, max_chars=240)
        return {
            "ok": True,
            "path": str(target),
            "name": target.name,
            "size": len(raw),
            "preview": preview,
        }
    except OSError as e:  # noqa: BLE001
        logging.warning(f"Couldn't save file {name}: {e}")
        return {"ok": False, "error": str(e)}


def list_files() -> list:
    """Return filenames (sorted) of all received files."""
    if not _FILES_DIR.exists():
        return []
    try:
        return sorted(
            p.name for p in _FILES_DIR.iterdir() if p.is_file()
        )
    except OSError:  # noqa: BLE001
        return []


def file_path(filename: str) -> Path | None:
    """Return the absolute Path for a stored file, or None if it doesn't exist."""
    name = _safe_filename(filename)
    p = _FILES_DIR / name
    if p.is_file():
        return p
    # Also try exact match (in case the file was saved under a non-normalized name).
    p2 = _FILES_DIR / filename
    return p2 if p2.is_file() else None


# ──────────────────────────────────────────────────────────────
# Active file pointer (like the default playlist)
# ──────────────────────────────────────────────────────────────
def set_active_file(filename: str) -> bool:
    """Persist which file is the active one the user is talking about."""
    try:
        _ACTIVE_FILE.write_text(
            json.dumps({"file": filename or ""}), encoding="utf-8"
        )
        return True
    except OSError as e:  # noqa: BLE001
        logging.warning(f"Couldn't save active file: {e}")
        return False


def get_active_file() -> str:
    """Return the active file's name, or '' if none set."""
    try:
        data = json.loads(_ACTIVE_FILE.read_text(encoding="utf-8"))
        name = data.get("file")
        if name and (_FILES_DIR / name).is_file():
            return name
    except (OSError, ValueError):  # noqa: BLE001
        pass
    return ""


# ──────────────────────────────────────────────────────────────
# Read content (for Peter to actually use the file)
# ──────────────────────────────────────────────────────────────
def _read_text_preview(raw: bytes, max_chars: int = 4000) -> str:
    """Best-effort plain-text extraction for a small blob of bytes."""
    try:
        text = raw.decode("utf-8", errors="ignore")
        text = " ".join(text.split())
        if len(text) > max_chars:
            return text[:max_chars] + "… [truncated]"
        return text
    except Exception:  # noqa: BLE001
        return ""


def read_file_content(filename: str, max_chars: int = 12000) -> str:
    """Extract readable text from a stored file for Peter to analyze.

    Supports common text/office/PDF formats. Returns a string that describes
    the file and its text, or '' if the file can't be found / is empty.
    """
    path = file_path(filename)
    if path is None:
        return ""
    try:
        raw = path.read_bytes()
        ext = path.suffix.lower()
        text = ""

        if ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(path)
                text = "\n".join(
                    (page.extract_text() or "") for page in reader.pages
                )
            except Exception as e:  # noqa: BLE001
                text = _read_text_preview(raw)
                logging.warning(f"PDF extract fallback for {filename}: {e}")
        elif ext in (".docx", ".doc"):
            try:
                import docx
                d = docx.Document(path)
                text = "\n".join(p.text for p in d.paragraphs)
                # Also pull table cells (resumes often use tables).
                for tbl in d.tables:
                    for row in tbl.rows:
                        cells = [c.text.strip() for c in row.cells]
                        if any(cells):
                            text += "\n" + " | ".join(cells)
            except Exception as e:  # noqa: BLE001
                text = _read_text_preview(raw)
                logging.warning(f"DOCX extract fallback for {filename}: {e}")
        else:
            # .txt, .md, .json, .csv, code, etc. — treat as UTF-8 text.
            text = raw.decode("utf-8", errors="ignore")
            text = " ".join(text.split())

        text = " ".join(text.split())
        if len(text) > max_chars:
            text = text[:max_chars] + "\n… [truncated]"
        return text
    except Exception as e:  # noqa: BLE001
        logging.error(f"read_file_content error for {filename}: {e}")
        return ""

